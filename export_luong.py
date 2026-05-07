"""
Xuất tài liệu "Luồng Hoạt Động Hệ Thống Phát Hiện Cá Mập Crypto" ra file Word.
Chạy: python export_luong.py
Yêu cầu: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_FILE = "Luong_Hoat_Dong_He_Thong_Ca_Map.docx"
DIAGRAM_IMG = r"C:\Users\nguye\.gemini\antigravity\brain\61fdaae5-e791-40db-a30e-6f97b97bec41\system_flowchart_1778027283089.png"

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Tô màu nền ô bảng."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    p.paragraph_format.space_after = Pt(4)
    return p

def add_code_block(doc, code: str):
    """Khối code với nền xám, font monospace."""
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)
        # nền xám nhạt
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F3F4F6')
        pPr.append(shd)
    doc.add_paragraph()  # khoảng cách

def add_info_table(doc, rows, headers, col_widths=None):
    """Thêm bảng thông tin với header màu."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1565C0')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_val in enumerate(row_data):
            row[i].text = cell_val
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ── Tạo tài liệu ───────────────────────────────────────────────────────────────

doc = Document()

# Lề trang
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Font mặc định
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ── Trang bìa ──────────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("🐋 LUỒNG HOẠT ĐỘNG HỆ THỐNG\nPHÁT HIỆN CÁ MẬP CRYPTO")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

doc.add_paragraph()
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("Hệ thống Big Data Real-time | Binance WebSocket → Kafka → PySpark → Streamlit")
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()
doc.add_page_break()

# ── 1. Sơ đồ tổng quan ────────────────────────────────────────────────────────
add_heading(doc, "1. Sơ Đồ Tổng Quan Luồng Dữ Liệu", level=1, color='0D47A1')
add_body(doc,
    "Hệ thống gồm 4 thành phần chính kết nối theo chuỗi pipeline: "
    "dữ liệu từ sàn Binance được thu thập qua WebSocket, đưa vào Kafka làm hàng đợi, "
    "PySpark xử lý và tổng hợp theo cửa sổ 5 giây, cuối cùng Streamlit hiển thị và phát cảnh báo."
)

if os.path.exists(DIAGRAM_IMG):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(DIAGRAM_IMG, width=Inches(6.2))
    cap = doc.add_paragraph("Hình 1. Sơ đồ luồng hoạt động hệ thống phát hiện cá mập")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
else:
    add_body(doc, "[Không tìm thấy file sơ đồ. Vui lòng kiểm tra đường dẫn DIAGRAM_IMG]",
             color='FF0000')

doc.add_paragraph()

# ── 2. Giai đoạn 1 ────────────────────────────────────────────────────────────
add_heading(doc, "2. Giai Đoạn 1 — Binance Gửi Dữ Liệu Thô", level=1, color='0D47A1')
add_heading(doc, "2.1 Hình dung thực tế", level=2)
add_body(doc,
    "Binance như một cái chợ khổng lồ. Mỗi khi có người mua/bán khớp lệnh, "
    "chợ phát loa thông báo tức thì qua WebSocket — không cần ai hỏi, dữ liệu tự đẩy về (Push model)."
)
add_body(doc, "Ví dụ thông báo liên tiếp trong 1 giây:", bold=True)
add_code_block(doc, """\
Có người vừa MUA  0.00126 BTC với giá $81,079.36
Có người vừa BÁN  0.00044 BTC với giá $81,078.00
Có người vừa MUA  0.00030 BTC với giá $81,079.36
... (hàng trăm thông báo mỗi giây)""")

add_heading(doc, "2.2 URL WebSocket (lấy từ tài liệu chính thức Binance)", level=2)
add_code_block(doc, """\
# Single stream (1 coin)
wss://stream.binance.com:9443/ws/btcusdt@trade

# Combined stream (nhiều coin — đang dùng trong hệ thống)
wss://stream.binance.com:9443/stream?streams=
    btcusdt@trade/ethusdt@trade/bnbusdt@trade/solusdt@trade/xrpusdt@trade""")

add_heading(doc, "2.3 Cấu trúc gói JSON Binance gửi về", level=2)
add_code_block(doc, """\
{
  "stream": "btcusdt@trade",
  "data": {
    "s": "BTCUSDT",        // Tên coin
    "p": "81079.36",       // Giá khớp lệnh (USD)
    "q": "0.00126",        // Số lượng BTC
    "m": false,            // false = MUA chủ động | true = BÁN chủ động
    "E": 1746508123000     // Thời gian (milliseconds)
  }
}""")

add_heading(doc, "2.4 Ý nghĩa trường is_buyer_maker (m)", level=2)
add_info_table(doc,
    rows=[
        ["false", "Người MUA đặt lệnh thị trường (Market Buy)", "Cộng vào buy_volume ✅"],
        ["true",  "Người BÁN đặt lệnh thị trường (Market Sell)", "Cộng vào sell_volume ❌"],
    ],
    headers=["Giá trị m", "Ý nghĩa", "Xử lý trong hệ thống"],
    col_widths=[1.0, 3.0, 2.5]
)
add_body(doc,
    "Đây là nền tảng để phân biệt áp lực mua/bán — cốt lõi của toàn bộ hệ thống phát hiện cá mập."
)

# ── 3. Giai đoạn 2 ────────────────────────────────────────────────────────────
add_heading(doc, "3. Giai Đoạn 2 — Kafka Làm Hàng Đợi Trung Gian", level=1, color='0D47A1')
add_heading(doc, "3.1 Hình dung thực tế", level=2)
add_body(doc,
    "Kafka như một bưu điện trung chuyển. Dữ liệu từ Binance được bỏ vào hòm thư (topic 'crypto_trades'). "
    "Dù Spark có đang bận xử lý, thư vẫn được giữ lại trong Kafka — không mất dữ liệu."
)

add_heading(doc, "3.2 Code gửi vào Kafka", level=2)
add_code_block(doc, """\
# binance_producer.py
producer = KafkaProducer(
    bootstrap_servers=['kafka:9092'],
    value_serializer=lambda v: json.dumps(v).encode('utf-8')
)

# Mỗi lệnh khớp → gửi ngay vào topic
producer.send('crypto_trades', {
    'symbol':        'BTCUSDT',
    'price':         81079.36,
    'quantity':      0.00126,
    'is_buyer_maker': False,    # BUY
    'timestamp':     1746508123000
})""")

add_heading(doc, "3.3 Lợi ích của Kafka", level=2)
add_info_table(doc,
    rows=[
        ["Không mất dữ liệu", "Nếu Spark chậm, Kafka giữ dữ liệu chờ"],
        ["Tách rời Producer/Consumer", "Binance Producer và Spark có thể chạy độc lập"],
        ["Chịu tải cao", "Xử lý hàng triệu message/giây"],
        ["Nhiều consumer", "Nhiều ứng dụng có thể đọc cùng 1 topic"],
    ],
    headers=["Lợi ích", "Giải thích"],
    col_widths=[2.5, 4.0]
)

# ── 4. Giai đoạn 3 ────────────────────────────────────────────────────────────
add_heading(doc, "4. Giai Đoạn 3 — PySpark Tổng Hợp Theo Cửa Sổ 5 Giây", level=1, color='0D47A1')
add_heading(doc, "4.1 Hình dung thực tế", level=2)
add_body(doc,
    "Spark như một kế toán viên. Cứ đúng 5 giây, anh ta cộng dồn tất cả các lệnh nhận được "
    "và viết 1 dòng tổng kết vào sổ (CSV). 847 lệnh nhỏ → 1 dòng tổng hợp."
)

add_heading(doc, "4.2 Ví dụ cụ thể", level=2)
add_code_block(doc, """\
Từ 07:10:20 → 07:10:25 (5 giây), nhận được 847 lệnh từ Kafka:
  - Tổng KL Mua (buy_volume):  29.0877 BTC
  - Tổng KL Bán (sell_volume):  1.0773 BTC
  - Giá trung bình (avg_price): $81,116.33
  - Giá cao nhất (high):        $81,120.00
  - Giá thấp nhất (low):        $81,110.00
→ Ghi 1 dòng vào realtime_crypto.csv""")

add_heading(doc, "4.3 Code PySpark tổng hợp", level=2)
add_code_block(doc, """\
# spark_streaming.py
windowed_df = parsed_df \\
    .withWatermark("timestamp", "5 seconds") \\
    .groupBy(
        window(col("timestamp"), "5 seconds"),   # Cửa sổ 5 giây
        col("symbol")
    ).agg(
        avg("price").alias("avg_price"),
        sum("quantity").alias("total_volume"),
        # is_buyer_maker=False → MUA chủ động
        sum(when(col("is_buyer_maker") == False, col("quantity"))).alias("buy_volume"),
        # is_buyer_maker=True  → BÁN chủ động
        sum(when(col("is_buyer_maker") == True,  col("quantity"))).alias("sell_volume"),
        max("price").alias("high"),
        min("price").alias("low"),
        first("price").alias("open"),
        last("price").alias("close")
    )""")

add_heading(doc, "4.4 Kết quả sau mỗi 5 giây", level=2)
add_info_table(doc,
    rows=[
        ["07:10:20", "BTCUSDT", "$81,116.33", "29.0877", "1.0773", "30.1650"],
        ["07:10:25", "BTCUSDT", "$81,145.14",  "4.2056", "2.8917",  "7.0972"],
        ["07:10:30", "BTCUSDT", "$81,130.00",  "0.1823", "0.0541",  "0.2364"],
    ],
    headers=["window_start", "symbol", "avg_price", "buy_volume", "sell_volume", "total_volume"],
    col_widths=[1.5, 1.0, 1.2, 1.1, 1.1, 1.2]
)

# ── 5. Giai đoạn 4 ────────────────────────────────────────────────────────────
add_heading(doc, "5. Giai Đoạn 4 — Streamlit Phát Hiện & Hiển Thị", level=1, color='0D47A1')
add_heading(doc, "5.1 Hình dung thực tế", level=2)
add_body(doc,
    "Dashboard như bảng điện tử ở sân bay — liên tục làm mới thông tin mỗi giây, "
    "tự động cảnh báo khi phát hiện sự kiện cá mập bất thường."
)

add_heading(doc, "5.2 Cơ chế phát hiện Cá Mập", level=2)
add_code_block(doc, """\
# dashboard.py
# Bước 1: Tính ngưỡng động
mean_vol  = coin_df['total_volume'].mean()    # KL trung bình 45 phiên (~3.75 phút)
threshold = mean_vol * whale_multiplier        # Nhân hệ số (mặc định 1×-5×)

# Bước 2: Lọc các phiên vượt ngưỡng
events = coin_df[coin_df['total_volume'] > threshold]

# Bước 3: Cảnh báo (chỉ 1 lần/phiên window mới)
if latest['total_volume'] > threshold:
    if st.session_state.get(alert_key) != latest_window_str:
        st.session_state[alert_key] = latest_window_str
        st.toast(f"🚨 PHÁT HIỆN CÁ MẬP! Giá: ${price} | KL: {vol}", icon="🐋")""")

add_heading(doc, "5.3 Ví dụ phát hiện thực tế", level=2)
add_code_block(doc, """\
Mean KL 45 phiên = 2.06 BTC
Hệ số nhạy      = 1.0×
Ngưỡng phát hiện = 2.06 × 1.0 = 2.06 BTC

Phiên 07:10:20: total_volume = 30.16 BTC > 2.06 → 🐋 CÁ MẬP! (Sell Whale - 96% bán)
Phiên 07:10:25: total_volume =  7.09 BTC > 2.06 → 🐋 CÁ MẬP! (Balanced - 59% mua)
Phiên 07:10:30: total_volume =  0.24 BTC < 2.06 → ✅ Bình thường""")

add_heading(doc, "5.4 Các chỉ số hiển thị trên Dashboard", level=2)
add_info_table(doc,
    rows=[
        ["Giá hiện tại",       "avg_price phiên mới nhất, xanh=tăng/đỏ=giảm so với open"],
        ["Lực Mua chủ động",   "buy_volume / total_volume × 100% → tỉ lệ % tiền mua"],
        ["Biến động (5s)",     "high - low trong phiên → biên độ dao động giá"],
        ["RSI (14)",           "Chỉ số sức mạnh tương đối tính trên 14 phiên gần nhất"],
        ["Mean KL (45 phiên)", "Khối lượng trung bình mỗi phiên 5s trong ~3.75 phút gần nhất"],
        ["Hệ số nhạy",         "Bội số nhân với Mean, điều chỉnh độ nhạy phát hiện"],
        ["Ngưỡng phát hiện",   "= Mean × Hệ số — phiên vượt ngưỡng này là cá mập"],
        ["Số phiên cá mập",    "Số phiên trong 45 phiên gần nhất được phân loại là cá mập"],
    ],
    headers=["Chỉ số", "Ý nghĩa"],
    col_widths=[2.0, 4.5]
)

# ── 6. Tóm tắt 3 điểm cốt lõi ────────────────────────────────────────────────
add_heading(doc, "6. Ba Điểm Cốt Lõi Cần Nhớ", level=1, color='0D47A1')
add_info_table(doc,
    rows=[
        ["1", "WebSocket = Push",    "Binance chủ động đẩy dữ liệu về mỗi lệnh khớp, không cần polling"],
        ["2", "Kafka = Buffer",      "Tách rời producer/consumer, đảm bảo không mất dữ liệu"],
        ["3", "Window 5s = Tổng hợp","Biến hàng trăm lệnh micro thành 1 tín hiệu vĩ mô để phát hiện cá mập"],
    ],
    headers=["#", "Công nghệ", "Vai trò"],
    col_widths=[0.4, 1.8, 4.3]
)

# ── Lưu file ───────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), OUTPUT_FILE)
doc.save(output_path)
print(f"[OK] Da xuat file: {output_path}")
