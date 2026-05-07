from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO ĐÁNH GIÁ DỰ ÁN BIG DATA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Hệ thống Phân tích Luồng tiền Crypto Real-time (Whale Tracker)\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Phần 1
    add_heading(doc, '1. Đánh giá tính phù hợp với môn học Big Data', level=1)
    doc.add_paragraph('Dự án này phản ánh kiến trúc của một hệ thống xử lý dữ liệu chuẩn ở cấp độ Doanh nghiệp (Enterprise Data Pipeline) - Kappa Architecture, đáp ứng hoàn hảo các tiêu chí của môn học Big Data:')
    
    add_heading(doc, '1.1. Đáp ứng Đặc trưng 4V của Big Data:', level=2)
    doc.add_paragraph('- Velocity (Tốc độ): Xử lý luồng dữ liệu Real-time từ WebSocket của Binance. Dữ liệu đổ về liên tục và tính toán ngay lập tức (Micro-batching 5 giây).', style='List Bullet')
    doc.add_paragraph('- Volume (Khối lượng): Giao dịch Crypto toàn cầu tạo ra hàng chục ngàn thao tác mỗi giây. Hệ thống phân tán đủ khả năng mở rộng để xử lý.', style='List Bullet')
    doc.add_paragraph('- Variety (Đa dạng): Chuyển đổi dữ liệu thô bán cấu trúc (JSON) thành dữ liệu có cấu trúc (Structured DataFrame).', style='List Bullet')
    doc.add_paragraph('- Veracity (Độ chân thực): Lọc nhiễu hàng nghìn giao dịch nhỏ, bóc tách hành vi định hướng (Taker Buy / Taker Sell) để phát hiện cá mập.', style='List Bullet')
    
    add_heading(doc, '1.2. Áp dụng Tech Stack chuẩn:', level=2)
    doc.add_paragraph('- Apache Kafka & Zookeeper: Đóng vai trò hệ thống Message Queue chịu tải cao, làm bộ đệm Ingestion cực tốt.', style='List Bullet')
    doc.add_paragraph('- Apache Spark (Streaming): Dùng tính toán phân tán, Windowing (5s), Watermarking để tổng hợp dữ liệu luồng.', style='List Bullet')
    doc.add_paragraph('- Docker: Ảo hóa Microservices, giả lập một cụm Cluster phân tán (Zookeeper, Kafka, Spark Master/Worker).', style='List Bullet')
    
    # Phần 2
    add_heading(doc, '2. Các thao tác chi tiết thực hiện dự án', level=1)
    p = doc.add_paragraph()
    p.add_run('Bước 1: Khởi tạo Môi trường và Infrastructure\n').bold = True
    p.add_run('- Xây dựng Dockerfile cho Streamlit.\n- Viết file docker-compose.yml để định nghĩa cụm 5 container: Zookeeper, Kafka, Spark-Master, Spark-Worker, Streamlit.\n- Cấu hình mạng lưới chung (crypto_net) và mount volumes để đồng bộ mã nguồn.')
    
    p = doc.add_paragraph()
    p.add_run('Bước 2: Phát triển Data Ingestion (Producer)\n').bold = True
    p.add_run('- Viết mã nguồn binance_producer.py kết nối với Websocket API của Binance.\n- Thu thập giao dịch của các đồng coin (BTC, ETH, SOL, BNB, XRP).\n- Đẩy dữ liệu liên tục dưới dạng JSON vào topic "crypto_trades" trên cụm Kafka.')
    
    p = doc.add_paragraph()
    p.add_run('Bước 3: Phát triển Data Processing (Spark Streaming)\n').bold = True
    p.add_run('- Cấu hình spark_streaming.py để pull dữ liệu từ Kafka.\n- Định nghĩa Schema để ép kiểu JSON sang DataFrame.\n- Áp dụng Window Function (5 giây) để tính tổng Khối lượng Mua/Bán chủ động, Giá cao, Giá thấp, Giá trung bình.\n- Xuất kết quả liên tục dưới dạng Dataframe Append vào file CSV (latest_crypto.csv).')
    
    p = doc.add_paragraph()
    p.add_run('Bước 4: Trực quan hóa dữ liệu (Streamlit Dashboard)\n').bold = True
    p.add_run('- Xây dựng dashboard.py bằng Streamlit và Plotly.\n- Thiết lập logic phát hiện "Cá mập": so sánh tổng lượng Buy/Sell trong khung 5 giây với ngưỡng thiết lập sẵn.\n- Sử dụng @st.fragment để auto-refresh biểu đồ và số liệu mỗi 2 giây mà không làm tải lại toàn bộ trang web.')
    
    p = doc.add_paragraph()
    p.add_run('Bước 5: Triển khai và Kiểm thử (Deployment)\n').bold = True
    p.add_run('- Dùng lệnh docker-compose up -d để khởi động hệ thống.\n- Xử lý các luồng dữ liệu trễ, tinh chỉnh các tham số kết nối giữa Kafka và Spark.')

    # Phần 3
    add_heading(doc, '3. Nhận xét tổng quan', level=1)
    doc.add_paragraph('✅ Điểm mạnh:', style='List Bullet')
    doc.add_paragraph('Hệ thống có độ trễ (latency) cực thấp, bắt kịp diễn biến thực tế của thị trường theo từng giây. Kiến trúc rành mạch, tính phân tán cao (decoupled) giúp dễ dàng cô lập và xử lý lỗi khi một thành phần bị quá tải.', style='List')
    doc.add_paragraph('⚠️ Hạn chế hiện tại:', style='List Bullet')
    doc.add_paragraph('Việc lưu trữ dữ liệu luồng vào file CSV cục bộ chưa thực sự tối ưu đối với khối lượng lớn. Trong thời gian dài, file có thể phình to hoặc việc đọc/ghi đồng thời giữa Spark và Pandas (Streamlit) có thể gây ra hiện tượng lock file.', style='List')

    # Phần 4
    add_heading(doc, '4. Định hướng phát triển tương lai', level=1)
    doc.add_paragraph('- Nâng cấp Storage Layer: Chuyển đổi từ việc lưu CSV sang một cơ sở dữ liệu chuỗi thời gian (Time-Series Database) chuyên dụng như InfluxDB hoặc NoSQL phân tán như Cassandra để đảm bảo khả năng ghi/đọc dữ liệu lớn tốc độ cao.', style='List Bullet')
    doc.add_paragraph('- Tích hợp Machine Learning (MLlib): Xây dựng mô hình dự đoán xu hướng giá (Price Prediction) hoặc phân loại hành vi bất thường (Anomaly Detection) tự động thay vì dùng các ngưỡng (thresholds) tĩnh cố định.', style='List Bullet')
    doc.add_paragraph('- Scale lên Cloud Architecture: Triển khai hệ thống lên AWS (sử dụng MSK cho Kafka, EMR cho Spark) thay vì chạy trên máy ảo local để khai thác năng lực điện toán mây.', style='List Bullet')
    doc.add_paragraph('- Hệ thống thông báo (Alerting System): Viết thêm một Service gửi cảnh báo Cá mập tự động qua Telegram Bot hoặc Discord Webhook cho người dùng.', style='List Bullet')

    doc.save('Bao_Cao_Tong_Quan_Du_An.docx')

if __name__ == "__main__":
    main()
